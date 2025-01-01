from flask import Blueprint, render_template, flash, request, jsonify,redirect,url_for,session
from flask import Flask, send_file, abort
from flask_login import login_required, current_user
from sqlalchemy.sql.functions import user
from werkzeug.exceptions import RequestURITooLarge
from googleapiclient.http import MediaIoBaseDownload
from .models import User_infor,File,Server,User,Class,UserClass
from . import db
from docx import Document
from todolist.control.Google import authenticate,word_to_pdf
from googleapiclient.errors import HttpError # type: ignore
from googleapiclient.http import MediaFileUpload,MediaIoBaseUpload
import os
from .check_format import check_report_format
from io import BytesIO
from werkzeug.utils import secure_filename
import win32com.client
from flask import send_from_directory
from datetime import datetime
views = Blueprint("views", __name__)
# https://drive.google.com/drive/u/1/folders/1LztIkqPOMfRwEMlpERaIkiIh_GngJp_9
# laasi id file của bạn 
PARENT_FOLDER_ID = "1LztIkqPOMfRwEMlpERaIkiIh_GngJp_9"
TYPE_FILE=['docx','pdf']
UPLOAD_FOLDER = 'file'
SERVER='Server'
AVATAR = 'avatar'
ALLOWED_EXTENSIONS = {'docx', 'pdf', 'txt'} 
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def checkformat(type):
    if type in TYPE_FILE:
        return 1
    else:
        return 0
@views.route("/", methods=["GET", "POST"])
@views.route("/home",methods=["GET","POST"])
def homedash():
      user_files = File.query.filter_by(user_id=current_user.id).all()    
      user_info = User_infor.query.filter_by(user_id=current_user.id).first()
      word_count = File.query.filter_by(user_id=current_user.id,filetype='.docx').count()
      pdf_count = File.query.filter_by(user_id=current_user.id,filetype='.pdf').count()
      pdf_convert = File.query.filter(File.user_id == current_user.id, File.link_convert != '').count()
      print(current_user.role)
      print(word_count,pdf_count,pdf_convert)
      print('doneeee')
      return render_template("home.html", user=current_user,user_info=user_info,user_files=user_files,word=word_count,pdf=pdf_count,convert=pdf_convert)

@views.route("/upload", methods=["GET", "POST"])
@login_required
def upload():
    if request.method == "POST":
        if "file" not in request.files:
            flash("No file part")
            return redirect(request.url)
        file = request.files["file"]
        # Truy cập các trường từ request.form
        title_name = request.form.get("title")  # Sử dụng request.form để lấy giá trị
        title_name1 = request.form.get("title1")
        title_name2 = request.form.get("title2")
        gv = request.form.get("giangvien")  # Sử dụng request.form để lấy giá trị
        lop = request.form.get("lop")  # Sử dụng request.form để lấy giá trị
        khoa = request.form.get("khoa")  # Sử dụng request.form để lấy giá trị
        loai_type = request.form.get("type")
        
        # In ra các giá trị
        print(gv, title_name,title_name1,title_name2, lop, khoa,loai_type,)
        # if file.filename.split(".")[1] not in TYPE_FILE:
        #     flash("Not help this type file",category="error")
        #     return redirect(request.url)
        print(file)
        print('day')
        print(file.filename.split('.')[0])
        title=file.filename.split('.')[0]
        file_extension = file.filename.rsplit('.', 1)[-1].lower()
        print(file_extension)
        TYPE_FILE=['docx','pdf']
        user_files = File.query.filter_by(user_id=current_user.id).all()    
        user_info = User_infor.query.filter_by(user_id=current_user.id).first()
        if file_extension in TYPE_FILE:
            if file.filename == "":
                flash("No selected file")
                return redirect(request.url)
            
            
            pdf_folder = os.path.join(r"E:\Todolist_Tutorial", SERVER)
            upload_folder = os.path.join(r"E:\Todolist_Tutorial", UPLOAD_FOLDER)

            os.makedirs(pdf_folder, exist_ok=True)
            os.makedirs(upload_folder, exist_ok=True)
                    
                
            filename = secure_filename(file.filename)
            if filename.endswith('.pdf'):
                local_file_path = os.path.join(pdf_folder, filename)
                file.save(local_file_path)
                print(f"File saved to: {local_file_path}")
                flash("File successfully uploaded locally")
                new_server = Server(
                        gvhd = gv,
                        filename = title,
                        title = title_name,
                        title1 = title_name1,
                        title2 = title_name2,
                        loai = loai_type,
                        lop = lop,
                        khoa =khoa,
                        link = local_file_path
                )
                db.session.add(new_server)
                db.session.commit()
                print('luu')
            else:
                local_file_path = os.path.join(upload_folder, filename)
                file.save(local_file_path)
                print(f"File saved to: {local_file_path}")
                flash("File successfully uploaded locally")
                print('no')
        
        
        
            service = authenticate()
            try:
                file_metadata={
                    'name': file.filename,
                    'parents':[PARENT_FOLDER_ID]
                }
                file_name, type = os.path.splitext(file.filename)
                file_content = BytesIO(file.read())  # Đọc và lưu tệp vào bộ nhớ
                

                # Tạo đối tượng MediaIoBaseUpload để tải tệp lên Google Drive
                media = MediaIoBaseUpload(file_content, mimetype='application/octet-stream')
                
                # Upload the file to Google Drive
                file = (
                    service.files().create(
                        body=file_metadata, 
                        media_body=media, 
                        fields='id,mimeType'
                    ).execute()
                )
                file_id = file.get('id')
                permission = {
                    'type': 'anyone',
                    'role': 'writer'
                }
                service.permissions().create(
                    fileId=file_id,
                    body=permission
                ).execute()
                file_Type = file.get('mimeType')
                print(file_id)
                print(file_name)
                print(type)
                # download_file1(file_id,file_name)
                if file_Type =='application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    print(f"https://docs.google.com/document/d/{file_id}/edit")
                    file_link = f"https://docs.google.com/document/d/{file_id}/edit"
                elif file_Type == 'application/pdf':
                    print(f"https://drive.google.com/file/d/{file_id}/view")
                    file_link = f"https://drive.google.com/file/d/{file_id}/view"
                else:
                    print(f"https://docs.google.com/presentation/d/{file_id}/edit")
                    file_link = f"https://docs.google.com/presentation/d/{file_id}/edit"
                
                new_file = File(
                    gvhd = gv,
                    title = title_name,
                    title1 = title_name1,
                    title2 = title_name2,
                    loai = loai_type,
                    lop = lop,
                    khoa =khoa,
                    filename=file_name,
                    filetype=type,
                    link=file_link,
                    user_id = current_user.id,
                    file_id = file_id,
                    file_word = local_file_path
                )
                db.session.add(new_file)
                db.session.commit()
                
                return redirect(request.url)
            except HttpError as error:
                print(f"An error occurred: {error}")
                file = None
                print('LMAF')
        else:
            return render_template("uploadfile.html", user=current_user,user_info=user_info,user_files=user_files)   
    user_files = File.query.filter_by(user_id=current_user.id).all()    
    user_info = User_infor.query.filter_by(user_id=current_user.id).first()

    return render_template("uploadfile.html", user=current_user,user_info=user_info,user_files=user_files)

@views.route("/convert",methods=["GET","POST"])
def convert():
      user_files = File.query.filter_by(user_id=current_user.id).all()    
      user_info = User_infor.query.filter_by(user_id=current_user.id).first()
      print('doneeee')
      return render_template("convert.html", user=current_user,user_info=user_info,user_files=user_files)

@views.route("/search",methods=["GET","POST"])
def search():
      user_files = File.query.filter_by(user_id=current_user.id).all()    
      user_info = User_infor.query.filter_by(user_id=current_user.id).first()
      print('search')
      return render_template("search.html", user=current_user,user_info=user_info,user_files=user_files)

@views.route("/view",methods=["GET","POST"])
def view():
      user_files = File.query.filter_by(user_id=current_user.id).all()    
      user_info = User_infor.query.filter_by(user_id=current_user.id).first()
      print('doneeee11')
      return render_template("viewword.html", user=current_user,user_info=user_info,user_files=user_files)

@views.route("/download_file_pdf", methods=["GET"])
@login_required
def download_file_pdf():
    file_path = request.args.get("file_pdf")
    print(file_path)
    try:
            # Kiểm tra xem file có tồn tại không
            with open(file_path, 'rb'):
                pass  # Nếu file mở được, tiếp tục xử lý
    except FileNotFoundError:
            abort(404, description="File không tồn tại!: {file_path}")  # Báo lỗi nếu không tìm thấy file

        # Trả file về cho người dùng
    return send_file(file_path, as_attachment=True)
  
@views.route("/convert_file", methods=["GET"])
@login_required
def convert_file():
    file_id = request.args.get("file_id")
    name = request.args.get("file_name")  
    file_path = request.args.get("file_path")
    file = File.query.filter_by(file_id=file_id).first()
    print(name)
    
    pdf_filename = name+".pdf"
    pdf_folder = os.path.join(r"E:\Todolist_Tutorial", SERVER,pdf_filename)
    pdf_file_path = os.path.join(r"D:\DHBK\GITHUB_PYTHON\web_framework\FLASK_PROJECT\Todolist_Tutorial", UPLOAD_FOLDER, pdf_filename)
    
    print(file_path)
    print(pdf_folder)
    print(pdf_file_path)
    print(file.filename)
    print(file.filetype)
    if file.link_convert:
        print("hehe")
    else:  
        print(file.link_convert)
        link = word_to_pdf(file_path,pdf_folder)
        print(link)
        new_server = Server(
                 gvhd = file.gvhd,
                 filename = file.filename,
                 title = file.title,
                 loai = file.loai,
                 lop = file.lop,
                 khoa =file.khoa,
                 link = link
            )
        db.session.add(new_server)
        db.session.commit()
        print('luu')
        # lnk,fileid= download_file(file_id,name)
        if file:
                # print(file.filename)
                print(f"File found: {file.link_convert}")
                file.link_convert = link
                # file.file_id_convert = fileid
                print(f"File found: {file.link_convert}")
                db.session.commit()

    user_files = File.query.filter_by(user_id=current_user.id).all()
    user_info = User_infor.query.filter_by(user_id=current_user.id).first()
    return render_template("convert.html", user=current_user,user_info=user_info,user_files=user_files)


@views.route('/fetch_file', methods=["GET", "POST"])
@login_required
def fetch_file():
    file_path = request.args.get("file_path")
    print(file_path)
    issues = check_report_format(file_path)
    issues_list = []
    for issue in issues:
        issues_list.append({
            'issue': issue['issue'],
            'recommend': issue['recommend']
        })
    # for issue in issues:
    #     print(issue)
    doc = Document(file_path)
    total_paragraphs = len(doc.paragraphs)
    content = []
    # Đọc từng phần của tài liệu
    for start in range(0, total_paragraphs, 100):
        end = min(start + 100, total_paragraphs)
        part_content = []
        for para in doc.paragraphs[start:end]:
            part_content.append(para.text)
        content.append("\n".join(part_content))
        # Xử lý phần này, ví dụ: lưu vào cơ sở dữ liệu, file, hoặc in ra
        # print("\n".join(part_content))  # Hoặc thực hiện các xử lý khác
        print(content)
        user_files = File.query.filter_by(user_id=current_user.id).all()    
        user_info = User_infor.query.filter_by(user_id=current_user.id).first()
    return render_template('viewword.html',user=current_user,user_info=user_info,user_files=user_files,context=content,issues=issues_list)

@views.route('/update-avatar', methods=['POST'])
def update_avatar():
    if 'avatar' not in request.files:
        return jsonify({'success': False, 'message': 'Không tìm thấy file ảnh!'}), 400

    avatar = request.files['avatar']
    if avatar.filename == '':
        return jsonify({'success': False, 'message': 'Tên file không hợp lệ!'}), 400

    # Kiểm tra định dạng file
    if not avatar.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
        return jsonify({'success': False, 'message': 'Định dạng file không hợp lệ! Chỉ hỗ trợ PNG, JPG, JPEG, GIF.'}), 400

    # Lưu file vào thư mục tĩnh
    # upload_folder = 'E:\Todolist_Tutorialstatic/uploads'
    
    upload_folder = os.path.join(r"E:\Todolist_Tutorial\todolist\static", AVATAR)
    os.makedirs(upload_folder, exist_ok=True)  # Tạo thư mục nếu chưa tồn tại
    file_path = os.path.join(upload_folder, avatar.filename)
    avatar.save(file_path)

    # Cập nhật URL ảnh đại diện trong database
    user_info = User_infor.query.filter_by(user_id=current_user.id).first()
    if user_info:
        user_info.avatar_url = f'avatar/{avatar.filename}'
        db.session.commit()
        return jsonify({'success': True, 'message': 'Avatar updated successfully!'})
    else:
        return jsonify({'success': False, 'message': 'User update Profile!'}), 404

# join class student
@views.route('/search_classes', methods=['GET', 'POST'])
def search_classes():
    user_files = File.query.filter_by(user_id=current_user.id).all()    
    user_info = User_infor.query.filter_by(user_id=current_user.id).first()
    print(current_user.role)
    print('doneeee')
    
    if request.method == 'POST':
        search_query = request.form['search_query']
        search_results = Class.query.filter(Class.name.ilike(f'%{search_query}%')).all()
        joined_classes = [cls.id for cls in current_user.classes]   
        print(joined_classes)
        return render_template('search_class.html', search_results=search_results,user=current_user,user_info=user_info,user_files=user_files,joined_classes=joined_classes)
    else:
        search_results = Class.query.all()
    joined_classes = [cls.id for cls in current_user.classes]   
    print(joined_classes)
    return render_template('search_class.html', search_results=[],user=current_user,user_info=user_info,user_files=user_files,joined_classes=joined_classes)

@views.route('/join_class/<int:class_id>', methods=['POST'])
def join_class(class_id):
    user_id = request.form['user_id']
    # user = User.query.get(session['user_id'])
    user = User.query.get(user_id)  # Lấy người dùng từ session
    class_to_join = Class.query.get(class_id)
    class_to_join.users.append(user)
    db.session.commit()
    return redirect('/search_classes')
@views.route('/leave_class/<int:class_id>', methods=['POST'])
def leave_class(class_id):
    user_id = request.form['user_id']
    # user = User.query.get(session['user_id'])  # Lấy người dùng từ session
    # class_to_leave = Class.query.get(class_id)  # Lấy lớp học từ ID

    # Kiểm tra xem người dùng có tham gia lớp này không
    existing_entry = db.session.query(UserClass).filter_by(user_id=user_id, class_id=class_id).first()

    if existing_entry:
        # Nếu có, xóa bản ghi trong bảng user_class (người dùng rời khỏi lớp)
        db.session.delete(existing_entry)
        db.session.commit()
        flash("You have successfully left the class!")  # Thông báo đã rời lớp
    else:
        # Nếu không có bản ghi, thông báo người dùng chưa tham gia lớp
        flash("You are not a member of this class!")  # Thông báo không tham gia lớp

    return redirect('/search_classes')  # Quay lại trang tìm kiếm lớp học


@views.route('/my_class',methods=['GET','POST'])
def my_class():
    user = User.query.get(current_user.id)
    user_files = File.query.filter_by(user_id=current_user.id).all()    
    user_info = User_infor.query.filter_by(user_id=current_user.id).first()
    user_classes = user.classes
    if user_classes:
        for cls in user_classes:
           print(f"Class ID: {cls.id}, Name: {cls.name}, Description: {cls.description}")
           return render_template('my_class.html',user=current_user,user_info=user_info,user_files=user_files,user_classes=user_classes)
    else:
        print("No classes found!")

    return render_template('my_class.html',user=current_user,user_info=user_info,user_files=user_files,user_classes=[])

    


# @views.route('/class_details/<int:class_id>', methods=['GET'])
# def class_details(class_id):
#     user_files = File.query.filter_by(user_id=current_user.id).all()    
#     user_info = User_infor.query.filter_by(user_id=current_user.id).first()
#     class_details = Class.query.get(class_id)  # Lấy lớp học theo class_id
#     if class_details:
#         students = class_details.users  # Lấy danh sách sinh viên tham gia lớp (từ bảng users)
#         return render_template('list_student_class.html',user=current_user,user_info=user_info,user_files=user_files,class_details=class_details, students=students)
#     else:
#         flash("Class not found.")
#         return redirect('/lecturer')  # Quay lại trang giảng viên nếu không tìm thấy lớp

# @views.route('/view_file_student/<int:id>',methods = ['GET'])
# def view_file_student(id):
#     # user_id = request.form['user_id']
#     # user_id = request.args.get('user_id')
#     # print(user_id)
#     print(id) 
#     user_info = User_infor.query.filter_by(user_id=current_user.id).first()
#     # = User.query.filter_by(user_id=user_id).first()
#     student =User_infor.query.filter_by(user_id=id).first()
#     print('k có à')
#     print(id)
#     print(student)
#     user_files = File.query.filter_by(user_id=id,status = 'approved').all()
#     print(user_files)
#     return render_template('lecturer_view_file_student.html',user=current_user,user_files=user_files,user_info=user_info,student=student)
# @views.route("/lecturer",methods=["GET","POST"])
# def lecturer():
#       user_files = File.query.filter_by(user_id=current_user.id).all()    
#       user_info = User_infor.query.filter_by(user_id=current_user.id).first()
#       classes = Class.query.all()
#       print(current_user.role)
#     #   files = File.query.filter_by(status='pending').all() 
#       return render_template("lecturer.html", user=current_user,user_info=user_info,user_files=user_files,classes=classes)
# @views.route('/create_class', methods=['POST'])
# def create_class():
#     # Lấy dữ liệu từ biểu mẫu
#     class_name = request.form['class_name']
#     description = request.form['description']
#     print(class_name,description)
#     # Kiểm tra tên lớp có tồn tại chưa
#     existing_class = Class.query.filter_by(name=class_name).first()
#     if existing_class:
#         flash('Tên lớp đã tồn tại!', category='error')
#         return redirect(url_for('views.lecturer'))
    
#     # Tạo lớp học mới
#     new_class = Class(name=class_name, description=description)
#     db.session.add(new_class)
#     db.session.commit()

#     user = current_user  # Giả sử bạn dùng Flask-Login để quản lý phiên đăng nhập
#     user.classes.append(new_class)  # Thêm lớp vào danh sách lớp của user
#     db.session.commit()

#     flash('Lớp học được tạo thành công!', category='success')
#     return redirect(url_for('views.lecturer'))
# @views.route('/delete_class/<int:class_id>', methods=['POST'])
# def delete_class(class_id):
#     cls = Class.query.get(class_id)
#     if not cls:
#         flash('Lớp học không tồn tại!', category='error')
#         return redirect(url_for('views.lecturer'))
#     # Xóa lớp học
#     db.session.delete(cls)
#     db.session.commit()
#     flash('Lớp học đã được xóa!', category='success')
#     return redirect(url_for('views.lecturer'))

# @views.route("/admin",methods=["GET","POST"])
# def admin():
#       user_files = File.query.filter_by(user_id=current_user.id).all()    
#       user_info = User_infor.query.filter_by(user_id=current_user.id).first()
#       print(current_user.role)
#       files = File.query.filter_by(status='pending').all() 
#       return render_template("admin.html", user=current_user,user_info=user_info,user_files=user_files,files=files)

# @views.route("/admin_manager",methods=["GET","POST"])
# def admin_manager():
#       user_files = File.query.filter_by(user_id=current_user.id).all()    
#       user_info = User_infor.query.filter_by(user_id=current_user.id).first()
#       users = User.query.filter(User.id != current_user.id).all()
#       print(users)
#       for i in users:
#           print(i.id,i.user_name,i.email,i.role)
       
#       print(current_user.role)
#       files = File.query.filter_by(status='pending').all() 
#       return render_template("user_management.html", user=current_user,user_info=user_info,user_files=user_files,files=files,users=users)

# @views.route('/approve_file/<int:file_id>', methods=['POST'])
# @login_required
# def approve_file(file_id):
#     # if current_user.role != 'admin':
#     #     flash('You are not authorized to approve files.', category='error')
#     #     return redirect(url_for('dashboard'))

#     file = File.query.get_or_404(file_id)
#     file.status = 'approved'  # Duyệt file
#     file.approval_time = datetime.utcnow()
#     db.session.commit()
#     flash('File approved successfully!', category='success')
#     return redirect(url_for('views.admin'))

# @views.route('/reject_file/<int:file_id>', methods=['POST'])
# @login_required
# def reject_file(file_id):
#     # if current_user.role != 'admin':
#     #     flash('You are not authorized to reject files.', category='error')
#     #     return redirect(url_for('dashboard'))

#     file = File.query.get_or_404(file_id)
#     file.status = 'rejected'  # Từ chối file
    
#     file.approval_time = None
#     db.session.commit()
#     flash('File rejected.', category='error')
#     return redirect(url_for('views.admin'))